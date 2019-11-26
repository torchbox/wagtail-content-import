from docx import Document

from .base import DocumentParser
from django.utils.html import format_html
from django.utils.safestring import mark_safe
from .tables import Table, Cell


class DocxParser(DocumentParser):
    """Default DocumentParser for docx taking a BytesIO docx file and converting upon parse() into a list of
    {'type': type, 'value': value} intermediate elements"""

    def __init__(self, document):
        self.document = Document(document)

    def generate_simple_tag(self, content, tag):
        return format_html('<{tag}>{content}</{tag}>', tag=tag, content=content) if content else ''

    def paragraph_to_html(self, paragraph, outer_tag='p'):
        """
        Compile a paragraph into a HTML string, optionally with semantic markup for styles.
        Returns a dictionary of the form:
        {
            'type': 'html',
            'value': html
        }
        """
        text_list = []
        for run in paragraph.runs:
            text = run.text
            if run.bold:
                text = self.generate_simple_tag(text, 'b')
            if run.italic:
                text = self.generate_simple_tag(text, 'em')
            text_list.append(text)

        content = mark_safe(''.join(text_list))

        return {
                'type': 'html',
                'value': self.generate_simple_tag(content, outer_tag)
                }

    def parse(self):
        """
        Parse the document and return a set of intermediate {'type': type, 'value': value} blocks that represent it.
        """
        blocks = []

        title = self.document.core_properties.title

        for paragraph in self.document.paragraphs:
            if paragraph.style.name == 'Heading 1':
                converted_block = {'type': 'heading', 'value': paragraph.text}
                if not title:
                    title = paragraph.text
            elif paragraph.style.name[:-1] == 'Heading ':
                converted_block = self.paragraph_to_html(paragraph, outer_tag='h'+paragraph.style.name[-1])
            else:
                converted_block = self.paragraph_to_html(paragraph)
            if converted_block['value']:
                blocks.append(converted_block)

        return {
            'title': title,
            'elements': blocks
        }
